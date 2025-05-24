#!/usr/bin/env python3
"""
Policy Analyzer Agent - Document Processing Pipeline

This module provides robust document processing capabilities for the Policy Analyzer Agent,
supporting multiple input formats and preparing documents for DORA compliance analysis.

Features:
- Multi-format support (PDF, DOCX, TXT, HTML)
- Document validation and integrity checks
- Text extraction with structure preservation
- Metadata extraction and enrichment
- Preprocessing for NLP analysis
- Error handling and recovery
- Performance monitoring
"""

import hashlib
import json
import logging
import mimetypes
import os
import re
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional, Tuple, Any, Union
from uuid import uuid4

# Document processing libraries
import fitz  # PyMuPDF for PDF processing
from docx import Document as DocxDocument
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph
import pytesseract
from PIL import Image
import io

# Text processing
import nltk
from nltk.tokenize import sent_tokenize, word_tokenize
from nltk.corpus import stopwords
import spacy

# Utilities
import pandas as pd
from tqdm import tqdm

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

@dataclass
class DocumentMetadata:
    """Document metadata structure"""
    document_id: str
    filename: str
    file_size: int
    file_type: str
    mime_type: str
    creation_date: Optional[datetime] = None
    modification_date: Optional[datetime] = None
    author: Optional[str] = None
    title: Optional[str] = None
    subject: Optional[str] = None
    keywords: List[str] = field(default_factory=list)
    language: Optional[str] = None
    page_count: Optional[int] = None
    word_count: Optional[int] = None
    character_count: Optional[int] = None
    processing_date: datetime = field(default_factory=datetime.now)
    md5_hash: Optional[str] = None

@dataclass
class DocumentSection:
    """Structured document section"""
    section_id: str
    title: str
    content: str
    level: int  # Hierarchy level (1=chapter, 2=section, 3=subsection, etc.)
    page_number: Optional[int] = None
    start_position: Optional[int] = None
    end_position: Optional[int] = None
    parent_section_id: Optional[str] = None
    subsections: List[str] = field(default_factory=list)
    tables: List[Dict[str, Any]] = field(default_factory=list)
    images: List[Dict[str, Any]] = field(default_factory=list)

@dataclass
class ProcessedDocument:
    """Complete processed document structure"""
    document_id: str
    metadata: DocumentMetadata
    raw_text: str
    structured_content: List[DocumentSection]
    tables: List[Dict[str, Any]] = field(default_factory=list)
    images: List[Dict[str, Any]] = field(default_factory=list)
    processing_stats: Dict[str, Any] = field(default_factory=dict)
    quality_score: float = 0.0
    
class DocumentValidationError(Exception):
    """Custom exception for document validation errors"""
    pass

class DocumentProcessingError(Exception):
    """Custom exception for document processing errors"""
    pass

class DocumentProcessor:
    """Main document processing pipeline"""
    
    # Supported file types
    SUPPORTED_EXTENSIONS = {'.pdf', '.docx', '.doc', '.txt', '.html', '.htm', '.rtf'}
    
    # Maximum file size (100MB)
    MAX_FILE_SIZE = 100 * 1024 * 1024
    
    # Text extraction settings
    MIN_TEXT_LENGTH = 100
    MAX_TEXT_LENGTH = 10_000_000
    
    def __init__(self, enable_ocr: bool = True, language: str = 'en'):
        """
        Initialize the document processor
        
        Args:
            enable_ocr: Enable OCR for scanned documents
            language: Language for text processing
        """
        self.enable_ocr = enable_ocr
        self.language = language
        
        # Initialize NLP components
        self._initialize_nlp()
        
        # Statistics tracking
        self.processing_stats = {
            'documents_processed': 0,
            'successful_extractions': 0,
            'failed_extractions': 0,
            'total_processing_time': 0.0,
            'average_processing_time': 0.0
        }
        
    def _initialize_nlp(self):
        """Initialize NLP libraries and models"""
        try:
            # Download required NLTK data
            nltk.download('punkt', quiet=True)
            nltk.download('stopwords', quiet=True)
            
            # Load spaCy model for advanced NLP
            try:
                self.nlp = spacy.load('en_core_web_sm')
            except OSError:
                logger.warning("spaCy English model not found. Basic processing only.")
                self.nlp = None
                
        except Exception as e:
            logger.error(f"Error initializing NLP: {e}")
            self.nlp = None
    
    def validate_document(self, file_path: Union[str, Path]) -> Tuple[bool, List[str]]:
        """
        Validate document before processing
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Tuple of (is_valid, list_of_issues)
        """
        file_path = Path(file_path)
        issues = []
        
        # Check if file exists
        if not file_path.exists():
            issues.append(f"File does not exist: {file_path}")
            return False, issues
        
        # Check file extension
        if file_path.suffix.lower() not in self.SUPPORTED_EXTENSIONS:
            issues.append(f"Unsupported file type: {file_path.suffix}")
        
        # Check file size
        file_size = file_path.stat().st_size
        if file_size > self.MAX_FILE_SIZE:
            issues.append(f"File too large: {file_size / (1024*1024):.1f}MB > {self.MAX_FILE_SIZE / (1024*1024):.1f}MB")
        
        # Check if file is empty
        if file_size == 0:
            issues.append("File is empty")
        
        # Check file permissions
        if not os.access(file_path, os.R_OK):
            issues.append("File is not readable")
        
        # Basic file integrity check
        try:
            with open(file_path, 'rb') as f:
                # Try to read first few bytes
                header = f.read(1024)
                if not header:
                    issues.append("File appears to be corrupted or empty")
        except Exception as e:
            issues.append(f"Error reading file: {e}")
        
        is_valid = len(issues) == 0
        return is_valid, issues
    
    def extract_metadata(self, file_path: Union[str, Path]) -> DocumentMetadata:
        """
        Extract metadata from document
        
        Args:
            file_path: Path to the document file
            
        Returns:
            DocumentMetadata object
        """
        file_path = Path(file_path)
        stat = file_path.stat()
        
        # Basic metadata
        metadata = DocumentMetadata(
            document_id=str(uuid4()),
            filename=file_path.name,
            file_size=stat.st_size,
            file_type=file_path.suffix.lower(),
            mime_type=mimetypes.guess_type(str(file_path))[0] or 'unknown',
            modification_date=datetime.fromtimestamp(stat.st_mtime),
            md5_hash=self._calculate_md5(file_path)
        )
        
        # Extract format-specific metadata
        try:
            if file_path.suffix.lower() == '.pdf':
                self._extract_pdf_metadata(file_path, metadata)
            elif file_path.suffix.lower() in ['.docx', '.doc']:
                self._extract_docx_metadata(file_path, metadata)
        except Exception as e:
            logger.warning(f"Could not extract advanced metadata: {e}")
        
        return metadata
    
    def _calculate_md5(self, file_path: Path) -> str:
        """Calculate MD5 hash of file"""
        hash_md5 = hashlib.md5()
        with open(file_path, "rb") as f:
            for chunk in iter(lambda: f.read(4096), b""):
                hash_md5.update(chunk)
        return hash_md5.hexdigest()
    
    def _extract_pdf_metadata(self, file_path: Path, metadata: DocumentMetadata):
        """Extract PDF-specific metadata"""
        try:
            doc = fitz.open(str(file_path))
            pdf_metadata = doc.metadata
            
            metadata.title = pdf_metadata.get('title')
            metadata.author = pdf_metadata.get('author')
            metadata.subject = pdf_metadata.get('subject')
            metadata.keywords = pdf_metadata.get('keywords', '').split(',') if pdf_metadata.get('keywords') else []
            metadata.page_count = len(doc)
            
            if pdf_metadata.get('creationDate'):
                # Parse PDF date format
                try:
                    date_str = pdf_metadata['creationDate']
                    # PDF date format: D:YYYYMMDDHHmmSSOHH'mm'
                    if date_str.startswith('D:'):
                        date_str = date_str[2:16]  # Take YYYYMMDDHHMMSS part
                        metadata.creation_date = datetime.strptime(date_str, '%Y%m%d%H%M%S')
                except Exception:
                    pass
            
            doc.close()
        except Exception as e:
            logger.warning(f"Error extracting PDF metadata: {e}")
    
    def _extract_docx_metadata(self, file_path: Path, metadata: DocumentMetadata):
        """Extract DOCX-specific metadata"""
        try:
            doc = DocxDocument(str(file_path))
            core_props = doc.core_properties
            
            metadata.title = core_props.title
            metadata.author = core_props.author
            metadata.subject = core_props.subject
            metadata.keywords = core_props.keywords.split(',') if core_props.keywords else []
            metadata.creation_date = core_props.created
            metadata.modification_date = core_props.modified
            
            # Count pages (approximate)
            metadata.page_count = len(doc.sections)
            
        except Exception as e:
            logger.warning(f"Error extracting DOCX metadata: {e}")
    
    def process_document(self, file_path: Union[str, Path], 
                        extract_tables: bool = True,
                        extract_images: bool = True) -> ProcessedDocument:
        """
        Main document processing method
        
        Args:
            file_path: Path to the document file
            extract_tables: Whether to extract tables
            extract_images: Whether to extract images
            
        Returns:
            ProcessedDocument object
        """
        start_time = datetime.now()
        file_path = Path(file_path)
        
        logger.info(f"Processing document: {file_path.name}")
        
        try:
            # Validate document
            is_valid, issues = self.validate_document(file_path)
            if not is_valid:
                raise DocumentValidationError(f"Document validation failed: {issues}")
            
            # Extract metadata
            metadata = self.extract_metadata(file_path)
            
            # Extract text based on file type
            if file_path.suffix.lower() == '.pdf':
                raw_text, tables, images = self._process_pdf(file_path, extract_tables, extract_images)
            elif file_path.suffix.lower() in ['.docx', '.doc']:
                raw_text, tables, images = self._process_docx(file_path, extract_tables, extract_images)
            elif file_path.suffix.lower() == '.txt':
                raw_text, tables, images = self._process_text(file_path)
            elif file_path.suffix.lower() in ['.html', '.htm']:
                raw_text, tables, images = self._process_html(file_path, extract_tables)
            else:
                raise DocumentProcessingError(f"Unsupported file type: {file_path.suffix}")
            
            # Validate extracted text
            if len(raw_text.strip()) < self.MIN_TEXT_LENGTH:
                raise DocumentProcessingError(f"Extracted text too short: {len(raw_text)} characters")
            
            # Structure the content
            structured_content = self._structure_content(raw_text)
            
            # Update metadata with text statistics
            metadata.word_count = len(raw_text.split())
            metadata.character_count = len(raw_text)
            
            # Calculate processing time
            processing_time = (datetime.now() - start_time).total_seconds()
            
            # Calculate quality score
            quality_score = self._calculate_quality_score(raw_text, structured_content, tables, images)
            
            # Create processed document
            processed_doc = ProcessedDocument(
                document_id=metadata.document_id,
                metadata=metadata,
                raw_text=raw_text,
                structured_content=structured_content,
                tables=tables,
                images=images,
                processing_stats={
                    'processing_time_seconds': processing_time,
                    'text_extraction_method': self._get_extraction_method(file_path),
                    'sections_identified': len(structured_content),
                    'tables_extracted': len(tables),
                    'images_extracted': len(images)
                },
                quality_score=quality_score
            )
            
            # Update statistics
            self.processing_stats['documents_processed'] += 1
            self.processing_stats['successful_extractions'] += 1
            self.processing_stats['total_processing_time'] += processing_time
            self.processing_stats['average_processing_time'] = (
                self.processing_stats['total_processing_time'] / 
                self.processing_stats['documents_processed']
            )
            
            logger.info(f"Successfully processed {file_path.name} in {processing_time:.2f}s "
                       f"(Quality: {quality_score:.2f})")
            
            return processed_doc
            
        except Exception as e:
            self.processing_stats['failed_extractions'] += 1
            logger.error(f"Error processing document {file_path.name}: {e}")
            raise DocumentProcessingError(f"Failed to process document: {e}")
    
    def _process_pdf(self, file_path: Path, extract_tables: bool, extract_images: bool) -> Tuple[str, List[Dict], List[Dict]]:
        """Process PDF document"""
        text_content = []
        tables = []
        images = []
        
        try:
            doc = fitz.open(str(file_path))
            
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                
                # Extract text
                text = page.get_text()
                if text.strip():
                    text_content.append(f"--- Page {page_num + 1} ---\n{text}")
                
                # Extract tables if requested
                if extract_tables:
                    page_tables = self._extract_pdf_tables(page, page_num)
                    tables.extend(page_tables)
                
                # Extract images if requested
                if extract_images:
                    page_images = self._extract_pdf_images(page, page_num)
                    images.extend(page_images)
                
                # OCR for pages with minimal text
                if self.enable_ocr and len(text.strip()) < 100:
                    ocr_text = self._ocr_page(page, page_num)
                    if ocr_text:
                        text_content.append(f"--- Page {page_num + 1} (OCR) ---\n{ocr_text}")
            
            doc.close()
            
        except Exception as e:
            raise DocumentProcessingError(f"Error processing PDF: {e}")
        
        raw_text = "\n\n".join(text_content)
        return raw_text, tables, images
    
    def _process_docx(self, file_path: Path, extract_tables: bool, extract_images: bool) -> Tuple[str, List[Dict], List[Dict]]:
        """Process DOCX document"""
        text_content = []
        tables = []
        images = []
        
        try:
            doc = DocxDocument(str(file_path))
            
            # Process document elements in order
            for element in doc.element.body:
                if isinstance(element, CT_P):
                    # Paragraph
                    paragraph = Paragraph(element, doc)
                    if paragraph.text.strip():
                        text_content.append(paragraph.text)
                
                elif isinstance(element, CT_Tbl) and extract_tables:
                    # Table
                    table = Table(element, doc)
                    table_data = self._extract_docx_table(table)
                    if table_data:
                        tables.append(table_data)
                        # Add table as text too
                        table_text = self._table_to_text(table_data)
                        text_content.append(f"[TABLE]\n{table_text}\n[/TABLE]")
            
            # Extract images if requested
            if extract_images:
                doc_images = self._extract_docx_images(doc)
                images.extend(doc_images)
            
        except Exception as e:
            raise DocumentProcessingError(f"Error processing DOCX: {e}")
        
        raw_text = "\n\n".join(text_content)
        return raw_text, tables, images
    
    def _process_text(self, file_path: Path) -> Tuple[str, List[Dict], List[Dict]]:
        """Process plain text document"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                raw_text = f.read()
        except UnicodeDecodeError:
            # Try with different encoding
            with open(file_path, 'r', encoding='latin-1') as f:
                raw_text = f.read()
        
        return raw_text, [], []
    
    def _process_html(self, file_path: Path, extract_tables: bool) -> Tuple[str, List[Dict], List[Dict]]:
        """Process HTML document"""
        try:
            from bs4 import BeautifulSoup
            
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            soup = BeautifulSoup(content, 'html.parser')
            
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()
            
            # Extract text
            text = soup.get_text()
            
            # Clean up text
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            text = '\n'.join(chunk for chunk in chunks if chunk)
            
            # Extract tables if requested
            tables = []
            if extract_tables:
                for table in soup.find_all('table'):
                    table_data = self._extract_html_table(table)
                    if table_data:
                        tables.append(table_data)
            
            return text, tables, []
            
        except Exception as e:
            raise DocumentProcessingError(f"Error processing HTML: {e}")
    
    def _structure_content(self, raw_text: str) -> List[DocumentSection]:
        """Structure the raw text into logical sections"""
        sections = []
        
        # Split text into potential sections based on common patterns
        section_patterns = [
            r'^[A-Z][A-Z\s]{10,}$',  # ALL CAPS headings
            r'^\d+\.\s+[A-Z]',       # Numbered sections (1. Title)
            r'^[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*:$',  # Title Case with colon
            r'^\s*#+\s+',            # Markdown-style headers
        ]
        
        lines = raw_text.split('\n')
        current_section = None
        section_content = []
        section_counter = 0
        
        for i, line in enumerate(lines):
            line_stripped = line.strip()
            
            # Check if this line is likely a section header
            is_header = False
            header_level = 1
            
            for pattern in section_patterns:
                if re.match(pattern, line_stripped):
                    is_header = True
                    break
            
            # Also check for short lines that might be headers
            if (len(line_stripped) < 100 and 
                len(line_stripped) > 5 and 
                not line_stripped.endswith('.') and
                line_stripped[0].isupper()):
                is_header = True
                header_level = 2
            
            if is_header and section_content:
                # Save previous section
                if current_section:
                    current_section.content = '\n'.join(section_content).strip()
                    sections.append(current_section)
                
                # Start new section
                section_counter += 1
                current_section = DocumentSection(
                    section_id=f"section_{section_counter}",
                    title=line_stripped,
                    content="",
                    level=header_level,
                    start_position=i
                )
                section_content = []
            else:
                # Add to current section
                section_content.append(line)
        
        # Add final section
        if current_section and section_content:
            current_section.content = '\n'.join(section_content).strip()
            sections.append(current_section)
        
        # If no sections were identified, create one section with all content
        if not sections:
            sections.append(DocumentSection(
                section_id="section_1",
                title="Document Content",
                content=raw_text,
                level=1
            ))
        
        return sections
    
    def _extract_pdf_tables(self, page, page_num: int) -> List[Dict[str, Any]]:
        """Extract tables from PDF page"""
        tables = []
        try:
            # Use pymupdf table extraction
            page_tables = page.find_tables()
            for i, table in enumerate(page_tables):
                table_data = {
                    'table_id': f"pdf_table_{page_num}_{i}",
                    'page_number': page_num + 1,
                    'rows': table.extract(),
                    'bbox': table.bbox
                }
                tables.append(table_data)
        except Exception as e:
            logger.warning(f"Could not extract tables from page {page_num}: {e}")
        
        return tables
    
    def _extract_pdf_images(self, page, page_num: int) -> List[Dict[str, Any]]:
        """Extract images from PDF page"""
        images = []
        try:
            image_list = page.get_images()
            for i, img in enumerate(image_list):
                image_data = {
                    'image_id': f"pdf_image_{page_num}_{i}",
                    'page_number': page_num + 1,
                    'xref': img[0],
                    'width': img[2],
                    'height': img[3]
                }
                images.append(image_data)
        except Exception as e:
            logger.warning(f"Could not extract images from page {page_num}: {e}")
        
        return images
    
    def _ocr_page(self, page, page_num: int) -> str:
        """Perform OCR on PDF page"""
        try:
            # Convert page to image
            pix = page.get_pixmap()
            img_data = pix.tobytes("png")
            img = Image.open(io.BytesIO(img_data))
            
            # Perform OCR
            ocr_text = pytesseract.image_to_string(img, lang='eng')
            return ocr_text
        except Exception as e:
            logger.warning(f"OCR failed for page {page_num}: {e}")
            return ""
    
    def _extract_docx_table(self, table: Table) -> Dict[str, Any]:
        """Extract table data from DOCX table"""
        rows = []
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                row_data.append(cell.text.strip())
            rows.append(row_data)
        
        return {
            'table_id': f"docx_table_{len(rows)}",
            'rows': rows,
            'row_count': len(rows),
            'col_count': len(rows[0]) if rows else 0
        }
    
    def _extract_docx_images(self, doc: DocxDocument) -> List[Dict[str, Any]]:
        """Extract images from DOCX document"""
        images = []
        try:
            for i, rel in enumerate(doc.part.rels.values()):
                if "image" in rel.target_ref:
                    image_data = {
                        'image_id': f"docx_image_{i}",
                        'relationship_id': rel.rId,
                        'target': rel.target_ref
                    }
                    images.append(image_data)
        except Exception as e:
            logger.warning(f"Could not extract DOCX images: {e}")
        
        return images
    
    def _extract_html_table(self, table) -> Dict[str, Any]:
        """Extract table data from HTML table"""
        rows = []
        for row in table.find_all('tr'):
            row_data = []
            for cell in row.find_all(['td', 'th']):
                row_data.append(cell.get_text().strip())
            if row_data:  # Only add non-empty rows
                rows.append(row_data)
        
        return {
            'table_id': f"html_table_{len(rows)}",
            'rows': rows,
            'row_count': len(rows),
            'col_count': len(rows[0]) if rows else 0
        }
    
    def _table_to_text(self, table_data: Dict[str, Any]) -> str:
        """Convert table data to readable text"""
        if not table_data.get('rows'):
            return ""
        
        text_lines = []
        for row in table_data['rows']:
            text_lines.append(" | ".join(str(cell) for cell in row))
        
        return "\n".join(text_lines)
    
    def _calculate_quality_score(self, raw_text: str, sections: List[DocumentSection], 
                                tables: List[Dict], images: List[Dict]) -> float:
        """Calculate document processing quality score (0-1)"""
        score = 0.0
        
        # Text quality (40% of score)
        if len(raw_text) > self.MIN_TEXT_LENGTH:
            text_score = min(1.0, len(raw_text) / 10000)  # Normalize to reasonable length
            score += text_score * 0.4
        
        # Structure quality (30% of score)
        if sections:
            structure_score = min(1.0, len(sections) / 10)  # Reward good structure
            score += structure_score * 0.3
        
        # Content richness (20% of score)
        content_score = 0
        if tables:
            content_score += 0.5
        if images:
            content_score += 0.5
        score += content_score * 0.2
        
        # Text clarity (10% of score)
        # Simple heuristic: ratio of alphabetic characters
        if raw_text:
            alpha_ratio = sum(c.isalpha() for c in raw_text) / len(raw_text)
            score += alpha_ratio * 0.1
        
        return min(1.0, score)
    
    def _get_extraction_method(self, file_path: Path) -> str:
        """Get the extraction method used for the file type"""
        ext = file_path.suffix.lower()
        method_map = {
            '.pdf': 'PyMuPDF + OCR',
            '.docx': 'python-docx',
            '.doc': 'python-docx',
            '.txt': 'Direct text',
            '.html': 'BeautifulSoup',
            '.htm': 'BeautifulSoup'
        }
        return method_map.get(ext, 'Unknown')
    
    def get_processing_statistics(self) -> Dict[str, Any]:
        """Get processing statistics"""
        return self.processing_stats.copy()
    
    def reset_statistics(self):
        """Reset processing statistics"""
        self.processing_stats = {
            'documents_processed': 0,
            'successful_extractions': 0,
            'failed_extractions': 0,
            'total_processing_time': 0.0,
            'average_processing_time': 0.0
        }


def main():
    """Main function for testing document processor"""
    import argparse
    
    parser = argparse.ArgumentParser(description="Document Processor Test")
    parser.add_argument("file_path", help="Path to document to process")
    parser.add_argument("--no-tables", action="store_true", help="Skip table extraction")
    parser.add_argument("--no-images", action="store_true", help="Skip image extraction")
    parser.add_argument("--no-ocr", action="store_true", help="Disable OCR")
    parser.add_argument("--output", help="Output file for processed document JSON")
    
    args = parser.parse_args()
    
    # Initialize processor
    processor = DocumentProcessor(enable_ocr=not args.no_ocr)
    
    try:
        # Process document
        processed_doc = processor.process_document(
            args.file_path,
            extract_tables=not args.no_tables,
            extract_images=not args.no_images
        )
        
        print(f"Successfully processed: {processed_doc.metadata.filename}")
        print(f"Document ID: {processed_doc.document_id}")
        print(f"File size: {processed_doc.metadata.file_size:,} bytes")
        print(f"Text length: {len(processed_doc.raw_text):,} characters")
        print(f"Word count: {processed_doc.metadata.word_count:,}")
        print(f"Sections: {len(processed_doc.structured_content)}")
        print(f"Tables: {len(processed_doc.tables)}")
        print(f"Images: {len(processed_doc.images)}")
        print(f"Quality score: {processed_doc.quality_score:.2f}")
        print(f"Processing time: {processed_doc.processing_stats['processing_time_seconds']:.2f}s")
        
        # Save to output file if specified
        if args.output:
            output_data = {
                'document_id': processed_doc.document_id,
                'metadata': processed_doc.metadata.__dict__,
                'raw_text': processed_doc.raw_text,
                'structured_content': [section.__dict__ for section in processed_doc.structured_content],
                'tables': processed_doc.tables,
                'images': processed_doc.images,
                'processing_stats': processed_doc.processing_stats,
                'quality_score': processed_doc.quality_score
            }
            
            with open(args.output, 'w') as f:
                json.dump(output_data, f, indent=2, default=str)
            
            print(f"Processed document saved to: {args.output}")
        
        # Show first few sections
        print("\nFirst few sections:")
        for i, section in enumerate(processed_doc.structured_content[:3]):
            print(f"{i+1}. {section.title} (Level {section.level})")
            preview = section.content[:200] + "..." if len(section.content) > 200 else section.content
            print(f"   {preview}\n")
        
    except Exception as e:
        print(f"Error processing document: {e}")
        return 1
    
    return 0


if __name__ == "__main__":
    exit(main()) 